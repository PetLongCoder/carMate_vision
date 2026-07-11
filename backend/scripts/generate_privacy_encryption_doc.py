"""Generate 用户隐私数据加密实现说明.docx — 2026-07-11."""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = r"C:\Users\18151\Desktop\用户隐私数据加密实现说明.docx"

PRIMARY = RGBColor(0x1A, 0x56, 0xDB)
ACCENT = RGBColor(0x0E, 0x7A, 0x6B)
DARK = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name: str = "微软雅黑", size: int | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_spacer(doc: Document, pts: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


def add_section_heading(doc: Document, num: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{num}  ")
    set_run_font(r1, size=14, bold=True, color=PRIMARY)
    r2 = p.add_run(title)
    set_run_font(r2, size=14, bold=True, color=DARK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "1A56DB")
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_font(run, bold=True, color=WHITE, size=10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            run = table.rows[r_idx + 1].cells[c_idx].paragraphs[0].add_run(val)
            set_run_font(run, size=10)
    add_spacer(doc, 8)


def build() -> None:
    doc = Document()
    set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("用户隐私数据加密实现说明"), size=22, bold=True, color=PRIMARY)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta.add_run("CarMate · 2026-07-11 · 负责人：罗天赐"), size=10, color=MUTED)
    add_spacer(doc, 10)

    add_section_heading(doc, "一", "实现目标")
    add_bullet(doc, "users 表手机号、邮箱、微信 OpenID 入库前 AES-256-GCM 加密")
    add_bullet(doc, "密码继续使用 bcrypt 单向哈希（不可逆）")
    add_bullet(doc, "API 返回脱敏值（138****8000 / u***r@example.com）")
    add_bullet(doc, "注册/登录/绑定仍可用明文输入，后端自动加密与查库")

    add_section_heading(doc, "二", "算法与密钥")
    add_table(
        doc,
        ["项目", "方案"],
        [
            ["对称加密", "AES-256-GCM（cryptography）"],
            ["确定性 nonce", "同一明文生成同一密文，支持 UNIQUE 与按手机号查用户"],
            ["密码", "bcrypt（已有）"],
            ["开发密钥", "DATA_ENCRYPTION_KEY，写在 backend/.env.example，全队 pull 即用"],
        ],
    )

    add_section_heading(doc, "三", "主要文件")
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["backend/app/core/crypto.py", "encrypt/decrypt/mask"],
            ["backend/app/services/user_privacy_service.py", "查库、赋值、迁移"],
            ["backend/app/api/v1/auth.py", "注册登录绑定接入加密"],
            ["backend/app/api/v1/wechat.py", "微信 OpenID 加密"],
            ["backend/scripts/migrate_encrypt_user_pii.py", "批量加密历史明文"],
        ],
    )

    add_section_heading(doc, "四", "队友使用")
    add_bullet(doc, "git pull 后无需改密钥，.env.example 已含开发密钥")
    add_bullet(doc, "首次启动后端会自动扩容字段并加密历史用户")
    add_bullet(doc, "DataGrip 查看 users.phone 应为 enc:v1: 开头密文")

    add_section_heading(doc, "五", "验证")
    add_bullet(doc, "python scripts/test_privacy_crypto.py")
    add_bullet(doc, "admin/123456 登录，用户中心手机显示 139****9000")
    add_bullet(doc, "手机验证码登录仍正常")

    add_spacer(doc, 12)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("— CarMate Vision —"), size=9, color=MUTED)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
