"""
生成交警手势识别日志功能文档 (DOCX)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return table

# ═══════════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_heading('交警手势识别 — 日志功能文档', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('CarMate Vision 车载视觉系统')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('云数据库日志存储 · 历史记录查询 · 视频分组展示')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. 概述
# ═══════════════════════════════════════════════════════════════
heading('1. 概述', 1)
para('交警手势识别日志功能将所有识别请求的结果持久化到腾讯云 MySQL 数据库中，支持按类型、手势名筛选查询，并在前端历史记录页面按视频分组展示。')
para('')
para('核心特性：')
para('• 所有识别类型（图片/视频/流式视频/摄像头实时帧）均自动写入日志')
para('• 视频识别按 video_session_id (UUID) 分组，每段手势单独一条记录')
para('• 异步线程池写入，不阻塞推理主流程')
para('• 前端历史记录页面按视频卡片分组展示，含动作时间线和段标签')

# ═══════════════════════════════════════════════════════════════
# 2. 架构
# ═══════════════════════════════════════════════════════════════
heading('2. 系统架构', 1)

heading('2.1 整体架构', 2)
para('日志系统分为四层：')
code_block('''
  识别服务 (police_gesture_service.py)
    → 日志构建 (GestureLogEntry)
      → 异步写入 (police_gesture_logger.py)
        → 云数据库 (police_gesture_logs 表)
          → 查询 API (GET /api/police-gesture/logs)
            → 前端展示 (History.tsx)''')

heading('2.2 核心文件', 2)
make_table(
    ['文件', '作用'],
    [
        ['backend/app/models/db_models.py', 'PoliceGestureLog ORM 模型定义'],
        ['backend/app/services/police_gesture_logger.py', '日志写入服务 (异步线程池)'],
        ['backend/app/services/police_gesture_service.py', '识别推理 + 日志调用集成'],
        ['backend/app/api/v1/police_gesture.py', '日志查询 API + 识别 API 错误日志'],
        ['backend/app/api/v1/history.py', '统一历史记录 API'],
        ['backend/app/api/v1/stats.py', '仪表盘统计 (含手势识别计数)'],
        ['frontend/src/pages/History.tsx', '前端历史记录页面 (按视频分组)'],
        ['frontend/src/store/policeGestureStore.ts', '识别结果 Zustand 跨页面持久化'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 3. 数据库设计
# ═══════════════════════════════════════════════════════════════
heading('3. 数据库设计', 1)

heading('3.1 数据库连接', 2)
make_table(
    ['参数', '值'],
    [
        ['主机', 'bj-cdb-mcxp1yss.sql.tencentcdb.com'],
        ['端口', '23196'],
        ['数据库', 'carmate'],
        ['用户', 'zbl'],
        ['驱动', 'PyMySQL + SQLAlchemy 2.0'],
        ['字符集', 'utf8mb4'],
    ]
)

heading('3.2 日志表结构 (police_gesture_logs)', 2)
make_table(
    ['字段', '类型', '说明'],
    [
        ['id', 'INT (PK, AUTO_INCREMENT)', '自增主键'],
        ['recognition_type', 'VARCHAR(32), INDEX', '识别类型: image/video/video_stream/camera_stream'],
        ['video_session_id', 'VARCHAR(36), INDEX', '视频会话 UUID，同视频的段共享 (视频模式)'],
        ['filename', 'VARCHAR(255)', '上传文件名'],
        ['gesture', 'VARCHAR(32)', '识别出的手势名称 (停止/直行/左转/...)'],
        ['gesture_id', 'INT', '手势 ID (0=无手势, 1-8)'],
        ['confidence', 'FLOAT', '置信度 (0~1)'],
        ['inference_ms', 'FLOAT', '推理耗时 (毫秒)'],
        ['top5_json', 'TEXT', 'Top-5 候选结果 JSON'],
        ['frames_total', 'INT', '视频总帧数 (视频模式)'],
        ['frames_processed', 'INT', '实际处理帧数 (视频模式)'],
        ['video_fps', 'FLOAT', '视频帧率 (视频模式)'],
        ['video_duration', 'FLOAT', '视频时长/秒 (视频模式)'],
        ['segments_json', 'TEXT', '手势段 JSON: [{start, end, gesture, gestureId}]'],
        ['success', 'BOOL', '识别是否成功'],
        ['error_message', 'TEXT', '错误信息 (失败时)'],
        ['client_info', 'VARCHAR(512)', '客户端信息 (预留)'],
        ['created_at', 'DATETIME, INDEX', '记录创建时间'],
    ]
)

heading('3.3 DDL (建表语句)', 2)
code_block('''CREATE TABLE police_gesture_logs (
    id              INT AUTO_INCREMENT PRIMARY KEY,
    recognition_type VARCHAR(32)  NOT NULL,
    video_session_id VARCHAR(36)  NULL,
    filename         VARCHAR(255) NULL,
    gesture          VARCHAR(32)  NOT NULL COMMENT '手势名称',
    gesture_id       INT          NOT NULL COMMENT '手势ID(0-8)',
    confidence       FLOAT        NOT NULL COMMENT '置信度',
    inference_ms     FLOAT        NOT NULL COMMENT '推理耗时(ms)',
    top5_json        TEXT         NULL     COMMENT 'Top-5结果JSON',
    frames_total     INT          NULL     COMMENT '视频总帧数',
    frames_processed INT          NULL     COMMENT '实际处理帧数',
    video_fps        FLOAT        NULL     COMMENT '视频帧率',
    video_duration   FLOAT        NULL     COMMENT '视频时长(秒)',
    segments_json    TEXT         NULL     COMMENT '手势段JSON',
    success          BOOL         DEFAULT TRUE,
    error_message    TEXT         NULL,
    client_info      VARCHAR(512) NULL,
    created_at       DATETIME     DEFAULT CURRENT_TIMESTAMP,
    INDEX idx_type (recognition_type),
    INDEX idx_video_session (video_session_id),
    INDEX idx_created (created_at)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;''')

# ═══════════════════════════════════════════════════════════════
# 4. 日志写入流程
# ═══════════════════════════════════════════════════════════════
heading('4. 日志写入流程', 1)

heading('4.1 GestureLogEntry 数据结构', 2)
code_block('''@dataclass
class GestureLogEntry:
    recognition_type: str       # image / video / video_stream / camera_stream
    gesture: str                # 手势名称
    gesture_id: int             # 手势 ID
    confidence: float           # 置信度
    inference_ms: float         # 推理耗时 (ms)
    success: bool = True        # 是否成功
    filename: Optional[str]     # 文件名
    video_session_id: Optional[str]  # 视频 UUID
    top5_json: Optional[str]    # Top-5 JSON
    frames_total: Optional[int] # 总帧数
    frames_processed: Optional[int] # 已处理帧数
    video_fps: Optional[float]  # 视频帧率
    video_duration: Optional[float] # 视频时长
    segments_json: Optional[str] # 手势段 JSON
    error_message: Optional[str] # 错误信息''')

heading('4.2 各识别模式的写入策略', 2)
make_table(
    ['识别模式', '写入时机', '分组方式', '说明'],
    [
        ['image (图片)', '识别完成后', '无分组', '单条记录，含 Top-5'],
        ['video (视频)', '每个手势段一条', 'video_session_id', '同一视频段共享 UUID'],
        ['video_stream (流式)', '每个手势段一条', 'video_session_id', '同 video 模式'],
        ['camera_stream (摄像头)', '新段确认时', '无分组', '仅有效手势 (>0) 且段变化时写入'],
    ]
)

heading('4.3 异步写入机制', 2)
para('日志写入使用单线程 ThreadPoolExecutor，避免阻塞推理主流程：')
code_block('''# 定义线程池 (单线程, 保证写入顺序)
_executor = ThreadPoolExecutor(max_workers=1)

# 异步写入 (不阻塞)
def log_gesture_async(entry: GestureLogEntry):
    _executor.submit(_write_log, entry)

# 批量写入 (同一事务)
def log_gestures_batch_async(entries: list[GestureLogEntry]):
    _executor.submit(_write_batch, entries)''')
para('写入失败仅记录 warning 日志，不影响 API 返回结果 (非致命设计)。')

# ═══════════════════════════════════════════════════════════════
# 5. API 接口
# ═══════════════════════════════════════════════════════════════
heading('5. API 接口', 1)

heading('5.1 交警手势日志查询', 2)
code_block('''GET /api/police-gesture/logs?page=1&pageSize=10&recognition_type=video&gesture=停止''')
para('参数：')
make_table(
    ['参数', '类型', '必填', '说明'],
    [
        ['page', 'int', '否', '页码 (默认 1)'],
        ['pageSize', 'int', '否', '每页条数 (默认 10, 最大 100)'],
        ['recognition_type', 'string', '否', '筛选类型: image/video/video_stream/camera_stream'],
        ['gesture', 'string', '否', '手势名称筛选: 停止/直行/左转/...'],
    ]
)
para('响应示例：')
code_block('''{
  "code": 200,
  "data": {
    "list": [{
      "id": 1,
      "recognitionType": "video",
      "videoSessionId": "a1b2c3d4-...",
      "filename": "test.mp4",
      "gesture": "停止",
      "gestureId": 1,
      "confidence": 0.952,
      "inferenceMs": 3200.5,
      "framesTotal": 475,
      "framesProcessed": 475,
      "segments": [{"start": 2.2, "end": 3.2, "gesture": "停止", "gestureId": 1}],
      "success": true,
      "createdAt": "2026-07-09T15:30:00"
    }],
    "total": 25,
    "page": 1,
    "pageSize": 10,
    "totalPages": 3
  }
}''')

heading('5.2 统一历史记录 API', 2)
code_block('''GET /api/history?page=1&pageSize=10&type=gesture''')
para('返回所有识别类型的历史记录 (当前含交警手势的 police_gesture_logs 数据)。')

heading('5.3 仪表盘统计 API', 2)
code_block('''GET /api/stats/dashboard''')
para('返回包含手势识别计数在内的统计数据：')
code_block('''{
  "totalPlates": 0,
  "totalGestures": 25,
  "todayGestures": 8,
  "successGestures": 23,
  "totalAlerts": 0,
  "unreadAlerts": 0
}''')

# ═══════════════════════════════════════════════════════════════
# 6. 视频分组机制
# ═══════════════════════════════════════════════════════════════
heading('6. 视频分组机制', 1)
para('每个视频上传时生成一个 UUID (video_session_id)，该视频的所有手势段共享此 ID。')
code_block('''# 在 process_police_gesture_video() 中:
video_session_id = str(uuid.uuid4())

# 每个手势段单独写入, 共享同一 video_session_id:
for seg in segments:
    GestureLogEntry(
        gesture=seg["gesture"],
        video_session_id=video_session_id,  # ← 共享 UUID
        segments_json=json.dumps([seg]),
        ...
    )''')
para('前端根据 videoSessionId 将多条记录合并为一个视频卡片，展示完整的时间线和手势段列表。')

# ═══════════════════════════════════════════════════════════════
# 7. 前端展示
# ═══════════════════════════════════════════════════════════════
heading('7. 前端展示', 1)
para('历史记录页面 (History.tsx) 按以下规则展示：')
para('• 同一 videoSessionId 的段合并为一张视频卡片')
para('• 卡片显示文件名、视频时长、帧数、推理耗时')
para('• 卡片内显示彩色动作时间线')
para('• 每个手势段有独立标签 (手势名 + 置信度 + 起止时间)')
para('• 图片/摄像头识别结果独立显示')
para('• 支持按类型筛选 (全部/车牌识别/交警手势/车主手势)')

# ═══════════════════════════════════════════════════════════════
# 8. 环境配置
# ═══════════════════════════════════════════════════════════════
heading('8. 环境配置', 1)
make_table(
    ['环境变量', '默认值', '说明'],
    [
        ['DB_HOST', 'bj-cdb-mcxp1yss.sql.tencentcdb.com', '腾讯云 MySQL 主机'],
        ['DB_PORT', '23196', '数据库端口'],
        ['DB_USER', 'zbl', '数据库用户'],
        ['DB_PASSWORD', 'zbl123456', '数据库密码'],
        ['DB_NAME', 'carmate', '数据库名称'],
        ['CARMATE_DEVICE', 'auto', '推理设备: auto/cpu/cuda'],
    ]
)
para('')
para('表在服务启动时通过 SQLAlchemy Base.metadata.create_all() 自动创建，无需手动建表。')

# ═══════════════════════════════════════════════════════════════
# 9. 数据示例
# ═══════════════════════════════════════════════════════════════
heading('9. 数据示例', 1)

heading('9.1 图片识别日志', 2)
code_block('''| id | recognition_type | gesture | confidence | inference_ms | created_at          |
| 1  | image            | 停止    | 0.952      | 120.5        | 2026-07-09 15:30:00 |''')

heading('9.2 视频识别日志 (同一视频 3 个段)', 2)
code_block('''| id | video_session_id                      | gesture | confidence | segments_json               |
| 2  | a1b2c3d4-e5f6-... | 停止    | 0.982      | [{"start":2.2,"end":3.2}]   |
| 3  | a1b2c3d4-e5f6-... | 直行    | 0.982      | [{"start":5.3,"end":7.0}]   |
| 4  | a1b2c3d4-e5f6-... | 左转    | 0.982      | [{"start":9.6,"end":11.1}]  |''')

# ═══════════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), '交警手势识别日志功能文档.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')
